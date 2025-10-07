"""
Office Document Analyzer Implementation
Analyzes Microsoft Office documents (Word, Excel, PowerPoint) to extract metadata.
"""

import logging
from pathlib import Path
from typing import Optional

from ..models.document_metadata import (
    DocumentMetadata, DocumentType, DocumentFormat, OfficeProperties,
    DocumentAuthor, DocumentProperties, TimestampInfo, ConfidenceLevel,
    DocumentClassification
)


class OfficeDocumentAnalyzerImpl:
    """
    Implementation of Office document analysis.

    Extracts metadata from Word, Excel, and PowerPoint documents.
    """

    def __init__(self):
        """Initialize the Office document analyzer."""
        self.logger = logging.getLogger(__name__)

    def analyze_docx(self, file_path: Path) -> DocumentMetadata:
        """
        Analyze a Word document and extract metadata.

        Args:
            file_path: Path to DOCX file

        Returns:
            DocumentMetadata with Word-specific information
        """
        from docx import Document

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        metadata = DocumentMetadata(
            file_path=file_path,
            file_name=file_path.name,
            file_size=file_path.stat().st_size,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_extension="docx"
        )

        try:
            doc = Document(str(file_path))

            # Extract core properties
            if doc.core_properties:
                metadata.author_info = self._extract_author_from_core_props(doc.core_properties)
                metadata.properties = self._extract_properties_from_core_props(doc.core_properties)
                metadata.timestamps = self._extract_timestamps_from_core_props(doc.core_properties)

            # Extract Word-specific properties
            metadata.office_properties = OfficeProperties()
            metadata.office_properties.application = "Microsoft Word"

            # Count paragraphs and estimate word count
            metadata.office_properties.paragraph_count = len(doc.paragraphs)
            metadata.office_properties.word_count = sum(
                len(paragraph.text.split()) for paragraph in doc.paragraphs
            )

            # Classification
            metadata.classification = DocumentClassification(
                document_type=DocumentType.OFFICE_DOCUMENT,
                document_format=DocumentFormat.DOCX,
                type_confidence=ConfidenceLevel.HIGH,
                format_confidence=ConfidenceLevel.HIGH,
                type_reason="Word document format",
                format_reason="Valid DOCX file"
            )

        except Exception as e:
            self.logger.error(f"Error analyzing DOCX {file_path}: {e}")
            metadata.extraction_errors.append(f"DOCX analysis error: {str(e)}")

        return metadata

    def analyze_xlsx(self, file_path: Path) -> DocumentMetadata:
        """
        Analyze an Excel document and extract metadata.

        Args:
            file_path: Path to XLSX file

        Returns:
            DocumentMetadata with Excel-specific information
        """
        from openpyxl import load_workbook

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        metadata = DocumentMetadata(
            file_path=file_path,
            file_name=file_path.name,
            file_size=file_path.stat().st_size,
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            file_extension="xlsx"
        )

        try:
            wb = load_workbook(str(file_path), read_only=True, data_only=True)

            # Extract core properties
            if wb.properties:
                metadata.author_info = self._extract_author_from_workbook_props(wb.properties)
                metadata.properties = self._extract_properties_from_workbook_props(wb.properties)
                metadata.timestamps = self._extract_timestamps_from_workbook_props(wb.properties)

            # Extract Excel-specific properties
            metadata.office_properties = OfficeProperties()
            metadata.office_properties.application = "Microsoft Excel"
            metadata.office_properties.sheet_count = len(wb.sheetnames)
            metadata.office_properties.sheet_names = wb.sheetnames

            # Classification
            metadata.classification = DocumentClassification(
                document_type=DocumentType.OFFICE_DOCUMENT,
                document_format=DocumentFormat.XLSX,
                type_confidence=ConfidenceLevel.HIGH,
                format_confidence=ConfidenceLevel.HIGH,
                type_reason="Excel document format",
                format_reason="Valid XLSX file"
            )

            wb.close()

        except Exception as e:
            self.logger.error(f"Error analyzing XLSX {file_path}: {e}")
            metadata.extraction_errors.append(f"XLSX analysis error: {str(e)}")

        return metadata

    def analyze_pptx(self, file_path: Path) -> DocumentMetadata:
        """
        Analyze a PowerPoint document and extract metadata.

        Args:
            file_path: Path to PPTX file

        Returns:
            DocumentMetadata with PowerPoint-specific information
        """
        from pptx import Presentation

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        metadata = DocumentMetadata(
            file_path=file_path,
            file_name=file_path.name,
            file_size=file_path.stat().st_size,
            mime_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            file_extension="pptx"
        )

        try:
            prs = Presentation(str(file_path))

            # Extract core properties
            if prs.core_properties:
                metadata.author_info = self._extract_author_from_core_props(prs.core_properties)
                metadata.properties = self._extract_properties_from_core_props(prs.core_properties)
                metadata.timestamps = self._extract_timestamps_from_core_props(prs.core_properties)

            # Extract PowerPoint-specific properties
            metadata.office_properties = OfficeProperties()
            metadata.office_properties.application = "Microsoft PowerPoint"
            metadata.office_properties.slide_count = len(prs.slides)

            # Classification
            metadata.classification = DocumentClassification(
                document_type=DocumentType.OFFICE_DOCUMENT,
                document_format=DocumentFormat.PPTX,
                type_confidence=ConfidenceLevel.HIGH,
                format_confidence=ConfidenceLevel.HIGH,
                type_reason="PowerPoint document format",
                format_reason="Valid PPTX file"
            )

        except Exception as e:
            self.logger.error(f"Error analyzing PPTX {file_path}: {e}")
            metadata.extraction_errors.append(f"PPTX analysis error: {str(e)}")

        return metadata

    def supports_format(self, file_extension: str) -> bool:
        """
        Check if format is supported.

        Args:
            file_extension: File extension (without dot)

        Returns:
            True if format is supported, False otherwise
        """
        supported = ["docx", "xlsx", "pptx", "odt", "ods", "odp"]
        return file_extension.lower() in supported

    def _extract_author_from_core_props(self, props) -> DocumentAuthor:
        """Extract author info from core properties (python-docx/python-pptx)."""
        author_info = DocumentAuthor()

        try:
            if hasattr(props, 'author') and props.author:
                author_info.author = str(props.author)
            if hasattr(props, 'last_modified_by') and props.last_modified_by:
                author_info.last_modified_by = str(props.last_modified_by)
            if hasattr(props, 'creator') and props.creator:
                author_info.creator = str(props.creator)
        except Exception as e:
            self.logger.warning(f"Error extracting author from core properties: {e}")

        return author_info

    def _extract_properties_from_core_props(self, props) -> DocumentProperties:
        """Extract document properties from core properties."""
        doc_props = DocumentProperties()

        try:
            if hasattr(props, 'title') and props.title:
                doc_props.title = str(props.title)
            if hasattr(props, 'subject') and props.subject:
                doc_props.subject = str(props.subject)
            if hasattr(props, 'keywords') and props.keywords:
                doc_props.keywords = [k.strip() for k in str(props.keywords).split(',') if k.strip()]
            if hasattr(props, 'category') and props.category:
                doc_props.category = str(props.category)
        except Exception as e:
            self.logger.warning(f"Error extracting properties from core properties: {e}")

        return doc_props

    def _extract_timestamps_from_core_props(self, props) -> TimestampInfo:
        """Extract timestamps from core properties."""
        timestamps = TimestampInfo()

        try:
            if hasattr(props, 'created') and props.created:
                timestamps.creation_date = props.created
            if hasattr(props, 'modified') and props.modified:
                timestamps.modification_date = props.modified
        except Exception as e:
            self.logger.warning(f"Error extracting timestamps from core properties: {e}")

        return timestamps

    def _extract_author_from_workbook_props(self, props) -> DocumentAuthor:
        """Extract author info from workbook properties (openpyxl)."""
        author_info = DocumentAuthor()

        try:
            if hasattr(props, 'creator') and props.creator:
                author_info.author = str(props.creator)
            if hasattr(props, 'lastModifiedBy') and props.lastModifiedBy:
                author_info.last_modified_by = str(props.lastModifiedBy)
        except Exception as e:
            self.logger.warning(f"Error extracting author from workbook properties: {e}")

        return author_info

    def _extract_properties_from_workbook_props(self, props) -> DocumentProperties:
        """Extract document properties from workbook properties."""
        doc_props = DocumentProperties()

        try:
            if hasattr(props, 'title') and props.title:
                doc_props.title = str(props.title)
            if hasattr(props, 'subject') and props.subject:
                doc_props.subject = str(props.subject)
            if hasattr(props, 'keywords') and props.keywords:
                doc_props.keywords = [k.strip() for k in str(props.keywords).split(',') if k.strip()]
            if hasattr(props, 'category') and props.category:
                doc_props.category = str(props.category)
        except Exception as e:
            self.logger.warning(f"Error extracting properties from workbook properties: {e}")

        return doc_props

    def _extract_timestamps_from_workbook_props(self, props) -> TimestampInfo:
        """Extract timestamps from workbook properties."""
        timestamps = TimestampInfo()

        try:
            if hasattr(props, 'created') and props.created:
                timestamps.creation_date = props.created
            if hasattr(props, 'modified') and props.modified:
                timestamps.modification_date = props.modified
        except Exception as e:
            self.logger.warning(f"Error extracting timestamps from workbook properties: {e}")

        return timestamps
