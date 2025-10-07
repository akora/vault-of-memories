"""
Integration tests for the complete document processing system.
Tests the interaction between all document processing components.
"""

import pytest
import tempfile
import time
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

from src.services.document_processor import DocumentProcessorImpl
from src.services.pdf_analyzer import PDFAnalyzerImpl
from src.services.office_document_analyzer import OfficeDocumentAnalyzerImpl
from src.services.ocr_detector import OCRDetectorImpl
from src.models.document_metadata import (
    DocumentType, DocumentFormat, ScanType, ConfidenceLevel
)


@pytest.mark.integration
class TestDocumentProcessingSystemIntegration:
    """Integration tests for complete document processing workflows."""

    def setup_method(self):
        """Set up test environment."""
        self.temp_dir = tempfile.mkdtemp()
        self.processor = DocumentProcessorImpl()
        self.pdf_analyzer = PDFAnalyzerImpl()
        self.office_analyzer = OfficeDocumentAnalyzerImpl()
        self.ocr_detector = OCRDetectorImpl()

    def teardown_method(self):
        """Clean up test environment."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def create_test_pdf(self, filename: str, page_count: int = 1, with_metadata: bool = False) -> Path:
        """Create a test PDF with optional metadata."""
        file_path = Path(self.temp_dir) / filename
        c = canvas.Canvas(str(file_path), pagesize=letter)

        # Add metadata if requested
        if with_metadata:
            c.setAuthor("Test Author")
            c.setTitle("Test Document Title")
            c.setSubject("Test Subject")

        for i in range(page_count):
            c.drawString(100, 750, f"Test Page {i + 1}")
            c.drawString(100, 700, "This is sample content for testing.")
            c.showPage()

        c.save()
        return file_path

    def create_test_docx(self, filename: str, with_metadata: bool = False) -> Path:
        """Create a test Word document."""
        file_path = Path(self.temp_dir) / filename
        doc = Document()

        if with_metadata:
            doc.core_properties.title = "Sample Document"
            doc.core_properties.author = "John Doe"
            doc.core_properties.subject = "Testing"

        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.save(str(file_path))
        return file_path

    def create_test_text_file(self, filename: str, content: str = None) -> Path:
        """Create a test text file."""
        file_path = Path(self.temp_dir) / filename
        if content is None:
            content = "This is a sample text file.\nWith multiple lines.\nFor testing purposes."
        file_path.write_text(content)
        return file_path

    def test_complete_brochure_processing_workflow(self):
        """Test complete workflow for brochure PDF (≤5 pages)."""
        file_path = self.create_test_pdf("brochure.pdf", page_count=3, with_metadata=True)

        # Process document
        metadata = self.processor.process_document(file_path)

        # Verify file information
        assert metadata.file_path == file_path
        assert metadata.mime_type == "application/pdf"
        assert metadata.file_size > 0

        # Verify PDF properties
        assert metadata.pdf_properties is not None
        assert metadata.pdf_properties.page_count == 3

        # Verify classification as brochure
        assert metadata.classification.document_type == DocumentType.BROCHURE
        assert metadata.classification.document_format == DocumentFormat.PDF
        assert metadata.classification.type_confidence == ConfidenceLevel.HIGH

        # Verify metadata extraction
        assert metadata.author_info.author == "Test Author"
        assert metadata.properties.title == "Test Document Title"
        assert metadata.properties.subject == "Test Subject"

        # Verify processing success
        assert metadata.is_successful()
        assert metadata.processing_time_ms > 0

    def test_complete_ebook_processing_workflow(self):
        """Test complete workflow for ebook PDF (>5 pages)."""
        file_path = self.create_test_pdf("ebook.pdf", page_count=50, with_metadata=True)

        metadata = self.processor.process_document(file_path)

        # Verify classification as ebook
        assert metadata.pdf_properties.page_count == 50
        assert metadata.classification.document_type == DocumentType.EBOOK
        assert metadata.pdf_properties.is_ebook()
        assert not metadata.pdf_properties.is_brochure()

    def test_boundary_page_count_classification(self):
        """Test page count boundary for brochure/ebook classification."""
        # Test at boundary (5 pages = brochure)
        boundary_file = self.create_test_pdf("boundary.pdf", page_count=5)
        metadata = self.processor.process_document(boundary_file)
        assert metadata.classification.document_type == DocumentType.BROCHURE

        # Test just over boundary (6 pages = ebook)
        over_boundary = self.create_test_pdf("over_boundary.pdf", page_count=6)
        metadata = self.processor.process_document(over_boundary)
        assert metadata.classification.document_type == DocumentType.EBOOK

    def test_word_document_processing_workflow(self):
        """Test complete workflow for Word documents."""
        file_path = self.create_test_docx("document.docx", with_metadata=True)

        metadata = self.processor.process_document(file_path)

        # Verify document type
        assert metadata.classification.document_type == DocumentType.OFFICE_DOCUMENT
        assert metadata.classification.document_format == DocumentFormat.DOCX

        # Verify metadata
        assert metadata.properties.title == "Sample Document"
        assert metadata.author_info.author == "John Doe"

        # Verify Office properties
        assert metadata.office_properties is not None
        assert metadata.office_properties.word_count > 0
        assert metadata.office_properties.paragraph_count >= 2

    def test_text_file_processing_workflow(self):
        """Test complete workflow for text files."""
        file_path = self.create_test_text_file("readme.txt")

        metadata = self.processor.process_document(file_path)

        assert metadata.classification.document_type == DocumentType.TEXT_FILE
        assert metadata.classification.document_format == DocumentFormat.TXT
        assert metadata.mime_type == "text/plain"
        assert 'line_count' in metadata.raw_metadata
        assert 'word_count' in metadata.raw_metadata

    def test_ocr_detection_workflow(self):
        """Test OCR/scan detection for digital-native PDF."""
        file_path = self.create_test_pdf("digital_native.pdf", page_count=1)

        metadata = self.processor.process_document(file_path)

        # Digital-native PDF should be detected
        assert metadata.ocr_detection.scan_type in [ScanType.DIGITAL_NATIVE, ScanType.UNKNOWN]

        # Should have text layer
        has_text = self.ocr_detector.has_text_layer(file_path)
        assert has_text is True

    def test_performance_benchmark_workflow(self):
        """Test performance of document processing."""
        # Create multiple test documents
        test_files = []

        # PDFs
        test_files.append(self.create_test_pdf("test1.pdf", page_count=3))
        test_files.append(self.create_test_pdf("test2.pdf", page_count=10))

        # Office documents
        test_files.append(self.create_test_docx("test3.docx"))

        # Text files
        test_files.append(self.create_test_text_file("test4.txt"))

        # Measure processing time
        start_time = time.time()
        results = []

        for file_path in test_files:
            metadata = self.processor.process_document(file_path)
            results.append(metadata)

        total_time = time.time() - start_time

        # Performance assertions
        assert total_time < 5.0  # Should complete in under 5 seconds
        assert len(results) == 4
        assert all(result.is_successful() for result in results)

        # Average processing time should be reasonable
        avg_processing_time = sum(r.processing_time_ms for r in results) / len(results)
        assert avg_processing_time < 2000  # Under 2 seconds per document

    def test_error_handling_workflow(self):
        """Test error handling for problematic documents."""
        # Create an empty file
        empty_file = Path(self.temp_dir) / "empty.pdf"
        empty_file.touch()

        # Processing should not crash
        metadata = self.processor.process_document(empty_file)

        # Should have errors but still return metadata object
        assert isinstance(metadata, object)
        # Empty/corrupted file will have extraction errors
        assert not metadata.is_successful()

    def test_multiple_document_types_workflow(self):
        """Test processing various document types."""
        documents = {
            'pdf_brochure': self.create_test_pdf("brochure.pdf", page_count=2),
            'pdf_ebook': self.create_test_pdf("ebook.pdf", page_count=20),
            'docx': self.create_test_docx("doc.docx", with_metadata=True),
            'txt': self.create_test_text_file("text.txt"),
        }

        results = {}
        for doc_type, file_path in documents.items():
            metadata = self.processor.process_document(file_path)
            results[doc_type] = metadata

        # Verify all processed successfully
        for doc_type, metadata in results.items():
            assert metadata.is_successful(), f"Failed to process {doc_type}"

        # Verify correct classifications
        assert results['pdf_brochure'].classification.document_type == DocumentType.BROCHURE
        assert results['pdf_ebook'].classification.document_type == DocumentType.EBOOK
        assert results['docx'].classification.document_type == DocumentType.OFFICE_DOCUMENT
        assert results['txt'].classification.document_type == DocumentType.TEXT_FILE

    def test_metadata_extraction_completeness(self):
        """Test completeness of metadata extraction."""
        file_path = self.create_test_pdf("complete.pdf", page_count=5, with_metadata=True)

        metadata = self.processor.process_document(file_path)

        # Verify all major components are present
        assert metadata.file_path is not None
        assert metadata.mime_type is not None
        assert metadata.pdf_properties is not None
        assert metadata.author_info is not None
        assert metadata.properties is not None
        assert metadata.timestamps is not None
        assert metadata.classification is not None
        assert metadata.ocr_detection is not None

        # Verify metadata can be serialized
        metadata_dict = metadata.to_dict()
        assert isinstance(metadata_dict, dict)
        assert 'file_path' in metadata_dict
        assert 'classification' in metadata_dict

    def test_page_count_accuracy_workflow(self):
        """Test page count accuracy across various sizes."""
        test_cases = [1, 3, 5, 6, 10, 25, 100]

        for expected_pages in test_cases:
            file_path = self.create_test_pdf(f"pages_{expected_pages}.pdf", page_count=expected_pages)
            metadata = self.processor.process_document(file_path)

            actual_pages = metadata.get_page_count()
            assert actual_pages == expected_pages, f"Expected {expected_pages} pages, got {actual_pages}"

    def test_timestamp_extraction_workflow(self):
        """Test timestamp extraction and fallback."""
        file_path = self.create_test_pdf("timestamped.pdf", with_metadata=True)

        metadata = self.processor.process_document(file_path)

        # Should have timestamps (either from metadata or file system)
        assert metadata.timestamps is not None

        # Should have a creation date (from file system at minimum)
        creation_date = metadata.get_creation_date()
        assert creation_date is not None

    def test_concurrent_processing_safety(self):
        """Test thread safety of document processing."""
        import threading

        file_path = self.create_test_pdf("concurrent.pdf", page_count=3, with_metadata=True)

        results = []
        errors = []

        def process_document_concurrently(thread_id):
            """Process document from multiple threads."""
            try:
                for _ in range(3):
                    metadata = self.processor.process_document(file_path)
                    results.append({
                        'thread_id': thread_id,
                        'success': metadata.is_successful(),
                        'page_count': metadata.get_page_count()
                    })
            except Exception as e:
                errors.append((thread_id, str(e)))

        # Start multiple threads
        threads = []
        for i in range(3):
            thread = threading.Thread(target=process_document_concurrently, args=(i,))
            threads.append(thread)
            thread.start()

        # Wait for completion
        for thread in threads:
            thread.join()

        # Verify results
        assert len(errors) == 0, f"Errors in concurrent processing: {errors}"
        assert len(results) == 9  # 3 threads × 3 processes each

        # All results should be consistent
        assert all(result['success'] for result in results)
        page_counts = [result['page_count'] for result in results]
        assert all(pc == 3 for pc in page_counts)

    def test_classification_confidence_levels(self):
        """Test confidence levels for various classifications."""
        # High confidence: PDF with clear page count
        pdf = self.create_test_pdf("confident.pdf", page_count=3)
        metadata = self.processor.process_document(pdf)
        assert metadata.classification.type_confidence == ConfidenceLevel.HIGH
        assert metadata.classification.format_confidence == ConfidenceLevel.HIGH

        # Office document should also have high confidence
        docx = self.create_test_docx("confident.docx")
        metadata = self.processor.process_document(docx)
        assert metadata.classification.type_confidence == ConfidenceLevel.HIGH

    def test_supported_formats_workflow(self):
        """Test all supported formats are handled."""
        supported = self.processor.get_supported_formats()

        assert "pdf" in supported
        assert "docx" in supported
        assert "xlsx" in supported
        assert "pptx" in supported
        assert "txt" in supported

    def test_can_process_check_workflow(self):
        """Test can_process file checking."""
        pdf_path = self.create_test_pdf("test.pdf")
        txt_path = self.create_test_text_file("test.txt")
        unsupported = Path(self.temp_dir) / "test.xyz"
        unsupported.touch()

        assert self.processor.can_process(pdf_path) is True
        assert self.processor.can_process(txt_path) is True
        assert self.processor.can_process(unsupported) is False

    def test_brochure_ebook_threshold_configurable(self):
        """Test that brochure/ebook threshold is configurable."""
        # Create analyzer with custom threshold
        custom_analyzer = PDFAnalyzerImpl(page_threshold=10)

        # 8 pages should be brochure with threshold of 10
        assert custom_analyzer.classify_pdf(8) == DocumentType.BROCHURE

        # 11 pages should be ebook
        assert custom_analyzer.classify_pdf(11) == DocumentType.EBOOK

    def test_password_protection_detection(self):
        """Test password protection detection."""
        # Regular PDF should not be password protected
        regular_pdf = self.create_test_pdf("regular.pdf")

        is_protected = self.pdf_analyzer.is_password_protected(regular_pdf)
        assert is_protected is False

    def test_large_document_handling(self):
        """Test handling of large documents."""
        # Create a large PDF
        large_pdf = self.create_test_pdf("large.pdf", page_count=200)

        start_time = time.time()
        metadata = self.processor.process_document(large_pdf)
        processing_time = time.time() - start_time

        # Should process successfully
        assert metadata.is_successful()
        assert metadata.pdf_properties.page_count == 200

        # Should complete in reasonable time
        assert processing_time < 10.0  # Under 10 seconds
